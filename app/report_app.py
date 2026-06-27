"""NMGone Report-/Auswertungsmodul ("Auswertungen").

Eigenstaendiges Programm wie die Kasse (eigenes Fenster / Taskleisten-Icon),
teilt sich die Datenbank mit NMGone (app/config.py -> ProgramData/NMGone).

Wichtig: Dieses Modul liest NUR aus tbl_bestellungen, tbl_bestellpositionen und
tbl_kunden_center. Es aendert NICHTS an den Kundendaten (keine Migration, kein
Schreiben) - der Kundenstamm gehoert ausschliesslich der Kunden App.

Drei Perspektiven (Verkaeufe / Kunden / Artikel). Pro Perspektive:
  - Suche : Zeitraum (akt. Monat / letzte 3-6-12 Monate / von-bis) + Freitextfilter
  - Ausgabe: frei waehl- und sortierbare Spalten (Dual-Listen-Dialog)
  - Export : Excel (immer), Word + PDF wenn die Bibliotheken installiert sind.

Umsaetze sind grundsaetzlich NETTO: apu * menge * (1 - rabatt_prozent/100).
"""
from __future__ import annotations

import calendar
from .i18n import T as _T
import json
import os
import sqlite3
import sys
from datetime import date, datetime
from pathlib import Path
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog

from .config import DB_PATH, ASSETS_DIR, OUTPUT_DIR, DEMO_SUFFIX
from . import theme

# ── Optik (zentrales Theme, gemeinsam mit NMGone / Kasse) ────────────────────
BG = theme.CARD
SHELL_BG = theme.BG
ACCENT = theme.PRIMARY
ACCENT2 = theme.ACCENT
NAV_SEL = theme.SELECT_BG
MUTED = theme.MUTED
CARD_BORDER = theme.BORDER
ZEBRA = "#F3F8FD"

FONT_H1 = (theme.FONT, 19, "bold")
FONT_SECTION = (theme.FONT, 11, "bold")
FONT_BTN = (theme.FONT, 10, "bold")
FONT_LIST = (theme.FONT, 11)
FONT_BASE = (theme.FONT, 10)


def _shade(hexcol, factor=0.86):
    """Dunklere Variante einer #rrggbb-Farbe (fuer Hover)."""
    try:
        h = hexcol.lstrip("#")
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        return f"#{int(r*factor):02x}{int(g*factor):02x}{int(b*factor):02x}"
    except Exception:
        return hexcol


def _eur(v) -> str:
    if v is None:
        return "—"
    try:
        return f"{float(v):,.2f} €".replace(",", "X").replace(".", ",").replace("X", ".")
    except (TypeError, ValueError):
        return str(v)


def _add_months(d: date, months: int) -> date:
    """Datum um -months Monate zurueck (months ist die Anzahl, positiv)."""
    total = (d.year * 12 + (d.month - 1)) - months
    y, m = divmod(total, 12)
    m += 1
    day = min(d.day, calendar.monthrange(y, m)[1])
    return date(y, m, day)


# ── Spalten-Definitionen je Perspektive ──────────────────────────────────────
# (key, label, typ)  typ: "text" | "int" | "eur"
COLS = {
    "verkaeufe": [
        ("kundennummer", "Kundennummer", "text"),
        ("apotheke", "Apotheke", "text"),
        ("inhaber", "Inhaber", "text"),
        ("strasse", "Straße", "text"),
        ("plz", "PLZ", "text"),
        ("ort", "Ort", "text"),
        ("email", "E-Mail", "text"),
        ("telefon", "Telefon", "text"),
        ("pk_zw", "PK/ZW", "text"),
        ("rechnungsart", "Rechnungsart", "text"),
        ("pzn", "PZN", "text"),
        ("artikel", "Artikel", "text"),
        ("menge", "Menge", "int"),
        ("umsatz_artikel", "Umsatz (Artikel)", "eur"),
        ("umsatz_kunde", "Umsatz (Kunde gesamt)", "eur"),
    ],
    "kunden": [
        ("kundennummer", "Kundennummer", "text"),
        ("apotheke", "Apotheke", "text"),
        ("inhaber", "Inhaber", "text"),
        ("strasse", "Straße", "text"),
        ("plz", "PLZ", "text"),
        ("ort", "Ort", "text"),
        ("email", "E-Mail", "text"),
        ("telefon", "Telefon", "text"),
        ("pk_zw", "PK/ZW", "text"),
        ("rechnungsart", "Rechnungsart", "text"),
        ("anzahl_artikel", "Artikel (versch.)", "int"),
        ("gesamt_menge", "Menge gesamt", "int"),
        ("gesamt_umsatz", "Umsatz gesamt", "eur"),
    ],
    "artikel": [
        ("pzn", "PZN", "text"),
        ("artikel", "Artikel", "text"),
        ("anzahl_kunden", "Kunden", "int"),
        ("gesamt_menge", "Menge gesamt", "int"),
        ("gesamt_umsatz", "Umsatz gesamt", "eur"),
    ],
}

# Default-Ausgabe je Perspektive (Reihenfolge = Default-Reihenfolge)
DEFAULTS = {
    "verkaeufe": ["kundennummer", "apotheke", "inhaber", "strasse", "plz", "ort",
                  "email", "telefon", "artikel", "menge", "umsatz_artikel",
                  "umsatz_kunde", "pk_zw", "rechnungsart"],
    "kunden": ["kundennummer", "apotheke", "inhaber", "plz", "ort", "telefon",
               "pk_zw", "rechnungsart", "anzahl_artikel", "gesamt_menge", "gesamt_umsatz"],
    "artikel": ["pzn", "artikel", "anzahl_kunden", "gesamt_menge", "gesamt_umsatz"],
}

PERSPEKTIVEN = [("verkaeufe", "Verkäufe"), ("kunden", "Kunden"),
                ("artikel", "Artikel"), ("frei", "Frei")]

# Netto-Umsatz-Ausdruck (SQL) - apu * menge * (1 - rabatt_prozent/100)
_NETTO = "p.menge * COALESCE(p.apu,0) * (1 - COALESCE(p.rabatt_prozent,0)/100.0)"

# ── Freie Auswertung ("Baukasten") ───────────────────────────────────────────
# Dimensionen = Gruppieren nach.  (key, label, sql, kunden_col|None)
# kunden_col gesetzt -> braucht tbl_kunden_center.<col>; fehlt sie -> NULL.
DIMENSIONS = [
    ("kundennummer", "Kundennummer", "b.kundennummer", None),
    ("apotheke",     "Apotheke",     "COALESCE(k.kundenname, b.apotheke)", "kundenname"),
    ("inhaber",      "Inhaber",      "k.inhaber", "inhaber"),
    ("ort",          "Ort",          "k.ort", "ort"),
    ("plz",          "PLZ",          "k.plz", "plz"),
    ("pk_zw",        "PK/ZW",        "k.kundentyp", "kundentyp"),
    ("rechnungsart", "Rechnungsart", "k.rechnungsart", "rechnungsart"),
    ("artikel",      "Artikel",      "p.artikelname", None),
    ("pzn",          "PZN",          "p.pzn", None),
    ("bestellart",   "Bestellart",   "COALESCE(p.bestellart,'Bestellung')", None),
    ("monat",        "Monat",        "strftime('%Y-%m', b.datum)", None),
    ("quartal",      "Quartal",      "strftime('%Y', b.datum) || '-Q' || ((CAST(strftime('%m', b.datum) AS INT)+2)/3)", None),
    ("jahr",         "Jahr",         "strftime('%Y', b.datum)", None),
]
# Kennzahlen = Werte.  (key, label, sql, typ)
MEASURES = [
    ("umsatz",           "Netto-Umsatz",       f"SUM({_NETTO})", "eur"),
    ("menge",            "Menge",              "SUM(p.menge)", "int"),
    ("anzahl_verkaeufe", "Anzahl Verkäufe",    "COUNT(DISTINCT b.id)", "int"),
    ("anzahl_artikel",   "Versch. Artikel",    "COUNT(DISTINCT p.pzn)", "int"),
    ("anzahl_kunden",    "Anzahl Kunden",      "COUNT(DISTINCT b.kundennummer)", "int"),
    ("avg_rabatt",       "Ø Rabatt %",         "AVG(COALESCE(p.rabatt_prozent,0))", "num"),
    ("avg_pos",          "Ø Umsatz/Position",  f"AVG({_NETTO})", "eur"),
]
# Filterbare Felder.  (key, label, sql, kind 'text'|'num', scope 'where'|'having', kunden_col|None)
FILTER_FIELDS = [
    ("kundennummer", "Kundennummer",     "b.kundennummer", "text", "where", None),
    ("apotheke",     "Apotheke",         "COALESCE(k.kundenname,b.apotheke)", "text", "where", "kundenname"),
    ("ort",          "Ort",              "k.ort", "text", "where", "ort"),
    ("plz",          "PLZ",              "k.plz", "text", "where", "plz"),
    ("pk_zw",        "PK/ZW",            "k.kundentyp", "text", "where", "kundentyp"),
    ("rechnungsart", "Rechnungsart",     "k.rechnungsart", "text", "where", "rechnungsart"),
    ("artikel",      "Artikel",          "p.artikelname", "text", "where", None),
    ("pzn",          "PZN",              "p.pzn", "text", "where", None),
    ("bestellart",   "Bestellart",       "COALESCE(p.bestellart,'Bestellung')", "text", "where", None),
    ("menge_pos",    "Menge (Position)", "p.menge", "num", "where", None),
    ("rabatt",       "Rabatt %",         "COALESCE(p.rabatt_prozent,0)", "num", "where", None),
    ("h_umsatz",     "Σ Netto-Umsatz (Gruppe)", f"SUM({_NETTO})", "num", "having", None),
    ("h_menge",      "Σ Menge (Gruppe)", "SUM(p.menge)", "num", "having", None),
]
DIM_MAP = {d[0]: d for d in DIMENSIONS}
MEA_MAP = {m[0]: m for m in MEASURES}
FILTER_MAP = {f[0]: f for f in FILTER_FIELDS}
OPS_TEXT = ["=", "≠", "enthält", "beginnt mit"]
OPS_NUM = ["=", "≠", ">", "<", "≥", "≤"]
_SQLOP = {"=": "=", "≠": "<>", ">": ">", "<": "<", "≥": ">=", "≤": "<="}


def _table_exists(con, name) -> bool:
    return con.execute(
        "SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (name,)
    ).fetchone() is not None


def _kunden_cols(con) -> set:
    if not _table_exists(con, "tbl_kunden_center"):
        return set()
    return {r[1] for r in con.execute("PRAGMA table_info(tbl_kunden_center)")}


def _k(expr_col: str, available: set, alias: str) -> str:
    """Liefert 'k.<col> AS <alias>' wenn die Spalte existiert, sonst 'NULL AS <alias>'.
    So bleibt die Abfrage robust, falls die Kunden App eine Spalte (noch) nicht hat.
    """
    return f"k.{expr_col} AS {alias}" if expr_col in available else f"NULL AS {alias}"


def _build_filter(spec: dict, kcols: set):
    """Baut aus {field, op, value} -> (sql_clause, params, scope) oder None.
    Whitelist ueber FILTER_MAP; Werte werden gebunden (kein Injection-Risiko).
    Filter auf nicht vorhandene Kunden-Spalten werden uebersprungen.
    """
    fdef = FILTER_MAP.get(spec.get("field"))
    if not fdef:
        return None
    key, label, expr, kind, scope, kcol = fdef
    if kcol and kcol not in kcols:
        return None  # Spalte (noch) nicht da -> Filter ignorieren
    op = spec.get("op") or "="
    val = spec.get("value", "")
    if kind == "num":
        try:
            num = float(str(val).replace(",", "."))
        except (TypeError, ValueError):
            return None
        return (f"{expr} {_SQLOP.get(op, '=')} ?", [num], scope)
    val = str(val)
    if op == "enthält":
        return (f"{expr} LIKE ?", [f"%{val}%"], scope)
    if op == "beginnt mit":
        return (f"{expr} LIKE ?", [f"{val}%"], scope)
    if op == "≠":
        return (f"COALESCE({expr},'') <> ?", [val], scope)
    return (f"{expr} = ?", [val], scope)


class BerichtPanel(tk.Frame):
    """Auswertungs-Oberflaeche. Laeuft als NMGone-Toplevel und als eigene .exe."""

    def __init__(self, master, db_path: Path = DB_PATH, on_close=None, nmgone_action=None):
        super().__init__(master, bg=SHELL_BG)
        self.db_path = str(db_path)
        self._on_close = on_close
        self._nmgone_action = nmgone_action
        self.perspektive = "verkaeufe"
        # gewaehlte Ausgabespalten je Perspektive (Reihenfolge bedeutsam)
        self.selected = {k: list(v) for k, v in DEFAULTS.items()}
        # Freie Auswertung: Dimensionen + Kennzahlen + Filter (Reihenfolge bedeutsam)
        self.builder = {"dims": ["ort"], "measures": ["umsatz", "menge"], "filters": []}
        self._result_cols: list[tuple] = []  # Spalten der zuletzt gerenderten Tabelle
        self._rows: list[dict] = []          # zuletzt ausgewertete Daten
        self._nav_buttons: dict[str, tk.Button] = {}

        self._apply_styles()
        self._build()
        self._select_perspektive("verkaeufe")

    # ── Stil + Bausteine ─────────────────────────────────────────────────────
    def _apply_styles(self):
        """Eigene, NUR fuer dieses Modul benannte ttk-Styles (klobbert NMGone nicht)."""
        style = ttk.Style(self)
        style.configure("Report.Treeview", rowheight=30, font=("Segoe UI", 11),
                        background="white", fieldbackground="white", borderwidth=0)
        style.configure("Report.Treeview.Heading", font=("Segoe UI", 10, "bold"),
                        background="#eaf1fa", foreground=ACCENT, padding=8, relief="flat")
        style.map("Report.Treeview", background=[("selected", NAV_SEL)],
                  foreground=[("selected", ACCENT)])
        style.map("Report.Treeview.Heading", background=[("active", "#dde9f7")])

    def _flat_btn(self, parent, text, cmd, bg=ACCENT, fg="white", big=False):
        b = tk.Button(parent, text=text, command=cmd, relief="flat", bd=0,
                      bg=bg, fg=fg, activebackground=_shade(bg), activeforeground=fg,
                      font=FONT_BTN, cursor="hand2",
                      padx=18 if big else 13, pady=9 if big else 7)
        b.bind("<Enter>", lambda e: b.configure(bg=_shade(bg)))
        b.bind("<Leave>", lambda e: b.configure(bg=bg))
        return b

    # ── Aufbau ───────────────────────────────────────────────────────────────
    def _build(self):
        self.columnconfigure(1, weight=1)
        self.rowconfigure(0, weight=1)

        # Linke Menueleiste (zentrale theme.Sidebar, einheitlich mit NMGone/Kasse)
        self.sidebar = theme.Sidebar(self, width=224, title="Auswertungen",
                                     subtitle="Verkäufe · Kunden · Artikel")
        self.sidebar.grid(row=0, column=0, sticky="ns")
        self._app_icon = theme.load_icon(ASSETS_DIR / "Report.ico", 56)
        if self._app_icon:
            self.sidebar.set_logo(self._app_icon)
        _icons = {"verkaeufe": "🧾", "kunden": "📇", "artikel": "🔍", "frei": "🧱"}
        for key, label in PERSPEKTIVEN:
            self.sidebar.add_item(key, _icons.get(key, "•"), label,
                                  lambda k=key: self._select_perspektive(k))

        nmb = tk.Button(self.sidebar.footer(), text="↩  NMGone öffnen", relief="flat",
                        bg=theme.SIDEBAR_ACTIVE, fg="white", activebackground="#1B5085",
                        activeforeground="white", bd=0, font=("Segoe UI", 10),
                        cursor="hand2", command=self._open_nmgone)
        nmb.pack(fill="x", padx=10, pady=(2, 18), ipady=8)

        # Rechter Inhalt
        content = tk.Frame(self, bg=SHELL_BG)
        content.grid(row=0, column=1, sticky="nsew")
        content.columnconfigure(0, weight=1)
        content.rowconfigure(2, weight=1)

        self.header = tk.Label(content, text="", bg=SHELL_BG, fg=ACCENT, font=FONT_H1)
        self.header.grid(row=0, column=0, sticky="w", padx=24, pady=(20, 10))

        # Such-Karte
        such = tk.Frame(content, bg=BG, highlightbackground=CARD_BORDER, highlightthickness=1)
        such.grid(row=1, column=0, sticky="ew", padx=24, pady=(0, 14))
        self._build_suche(such)

        # Ergebnis-Karte
        erg = tk.Frame(content, bg=BG, highlightbackground=CARD_BORDER, highlightthickness=1)
        erg.grid(row=2, column=0, sticky="nsew", padx=24, pady=(0, 20))
        erg.columnconfigure(0, weight=1)
        erg.rowconfigure(0, weight=1)
        self.tree = ttk.Treeview(erg, show="headings", selectmode="browse", style="Report.Treeview")
        self.tree.grid(row=0, column=0, sticky="nsew", padx=(12, 0), pady=12)
        self.tree.tag_configure("odd", background=ZEBRA)
        self.tree.tag_configure("even", background="white")
        vsb = ttk.Scrollbar(erg, orient="vertical", command=self.tree.yview)
        vsb.grid(row=0, column=1, sticky="ns", pady=12)
        hsb = ttk.Scrollbar(erg, orient="horizontal", command=self.tree.xview)
        hsb.grid(row=1, column=0, sticky="ew", padx=(12, 0))
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        self.status = tk.Label(erg, text="", bg=BG, fg=MUTED, font=("Segoe UI", 10))
        self.status.grid(row=2, column=0, columnspan=2, sticky="w", padx=14, pady=(0, 8))

    def _build_suche(self, parent):
        parent.columnconfigure(6, weight=1)
        tk.Label(parent, text="Zeitraum", bg=BG, fg=ACCENT,
                 font=FONT_SECTION).grid(row=0, column=0, sticky="w", padx=(16, 10), pady=(14, 8))
        self.zeitraum = tk.StringVar(value="3m")
        for i, (val, lab) in enumerate([("akt", "Akt. Monat"), ("3m", "3 Monate"),
                                        ("6m", "6 Monate"), ("12m", "12 Monate"),
                                        ("range", "Von–Bis")]):
            tk.Radiobutton(parent, text=lab, value=val, variable=self.zeitraum, bg=BG,
                           font=FONT_BASE, activebackground=BG, cursor="hand2",
                           command=self._toggle_range).grid(row=0, column=1 + i, sticky="w", padx=4, pady=(14, 8))

        # Von / Bis (nur bei "range")
        self.range_frame = tk.Frame(parent, bg=BG)
        self.range_frame.grid(row=1, column=0, columnspan=8, sticky="w", padx=16, pady=(0, 8))
        self._von, self._bis = self._make_date_inputs(self.range_frame)

        # Freitext-Filter + Buttons
        bottom = tk.Frame(parent, bg=BG)
        bottom.grid(row=2, column=0, columnspan=8, sticky="ew", padx=16, pady=(2, 14))
        bottom.columnconfigure(1, weight=1)
        tk.Label(bottom, text="Filter", bg=BG, fg=ACCENT,
                 font=FONT_SECTION).grid(row=0, column=0, sticky="w")
        self.filter_var = tk.StringVar()
        ent = tk.Entry(bottom, textvariable=self.filter_var, font=FONT_BASE)
        ent.grid(row=0, column=1, sticky="ew", padx=10, ipady=4)
        ent.bind("<Return>", lambda e: self.run_query())
        self._cfg_btn = self._flat_btn(bottom, "🔧 Spalten…", self._open_config,
                                       bg="#e8eef6", fg=ACCENT)
        self._cfg_btn.grid(row=0, column=2, padx=4)
        self._flat_btn(bottom, "🔎 Auswerten", self.run_query, bg=ACCENT, big=True).grid(
            row=0, column=3, padx=4)
        self._flat_btn(bottom, "⬇ Excel", lambda: self.export("excel"), bg="#11823b").grid(
            row=0, column=4, padx=4)
        self._flat_btn(bottom, "⬇ Word", lambda: self.export("word"), bg="#2b5797").grid(
            row=0, column=5, padx=4)
        self._flat_btn(bottom, "⬇ PDF", lambda: self.export("pdf"), bg="#a23b3b").grid(
            row=0, column=6, padx=4, sticky="w")
        self._toggle_range()

    def _make_date_inputs(self, parent):
        tk.Label(parent, text="Von:", bg=BG).pack(side="left")
        try:
            from tkcalendar import DateEntry
            von = DateEntry(parent, date_pattern="yyyy-mm-dd", width=12)
            von.pack(side="left", padx=(4, 12))
            tk.Label(parent, text="Bis:", bg=BG).pack(side="left")
            bis = DateEntry(parent, date_pattern="yyyy-mm-dd", width=12)
            bis.pack(side="left", padx=4)
        except Exception:
            von = tk.Entry(parent, width=12)
            von.insert(0, _add_months(date.today(), 3).isoformat())
            von.pack(side="left", padx=(4, 12))
            tk.Label(parent, text="Bis (JJJJ-MM-TT):", bg=BG).pack(side="left")
            bis = tk.Entry(parent, width=12)
            bis.insert(0, date.today().isoformat())
            bis.pack(side="left", padx=4)
        return von, bis

    def _toggle_range(self):
        state = "normal" if self.zeitraum.get() == "range" else "disabled"
        for w in self.range_frame.winfo_children():
            try:
                w.configure(state=state)
            except tk.TclError:
                pass

    # ── Perspektivenwechsel ──────────────────────────────────────────────────
    def _select_perspektive(self, key):
        self.perspektive = key
        self.sidebar.set_active(key)
        label = dict(PERSPEKTIVEN)[key]
        self.header.configure(text=_T('Auswertung · {p0}', p0=label))
        self._cfg_btn.configure(text="🧱 Bauen…" if key == "frei" else "🔧 Spalten…")
        self.run_query()

    def _open_config(self):
        if self.perspektive == "frei":
            self._open_builder_dialog()
        else:
            self._open_columns_dialog()

    # ── Datumsbereich aus Auswahl ────────────────────────────────────────────
    def _date_range(self):
        today = date.today()
        z = self.zeitraum.get()
        if z == "akt":
            return today.replace(day=1).isoformat(), today.isoformat()
        if z == "range":
            return self._read_date(self._von), self._read_date(self._bis)
        months = {"3m": 3, "6m": 6, "12m": 12}.get(z, 3)
        return _add_months(today, months).isoformat(), today.isoformat()

    @staticmethod
    def _read_date(widget) -> str:
        try:
            return widget.get_date().isoformat()      # DateEntry
        except Exception:
            return (widget.get() or "").strip()

    # ── Abfragen ─────────────────────────────────────────────────────────────
    def run_query(self):
        von, bis = self._date_range()
        try:
            con = sqlite3.connect(self.db_path)
            con.row_factory = sqlite3.Row
            if not (_table_exists(con, "tbl_bestellpositionen") and _table_exists(con, "tbl_bestellungen")):
                self._rows = []
                self._render([], "Noch keine Verkaufsdaten vorhanden.")
                return
            kcols = _kunden_cols(con)
            if self.perspektive == "verkaeufe":
                rows = self._q_verkaeufe(con, von, bis, kcols)
            elif self.perspektive == "kunden":
                rows = self._q_kunden(con, von, bis, kcols)
            elif self.perspektive == "artikel":
                rows = self._q_artikel(con, von, bis)
            else:  # frei
                rows, self._result_cols = self._q_frei(con, von, bis, kcols)
        except Exception as exc:
            self._render([], f"Fehler bei der Auswertung: {exc}")
            return
        finally:
            try:
                con.close()
            except Exception:
                pass

        # Freitextfilter (ueber alle gewaehlten Textspalten)
        flt = self.filter_var.get().strip().lower()
        if flt:
            keys = [k for k, _, _ in self._active_cols()]
            rows = [r for r in rows if any(flt in str(r.get(k, "") or "").lower() for k in keys)]
        self._rows = rows
        self._render(rows, f"{len(rows)} Zeilen · Zeitraum {von} – {bis}")

    def _kunden_select(self, kcols):
        return ", ".join([
            _k("inhaber", kcols, "inhaber"), _k("strasse", kcols, "strasse"),
            _k("plz", kcols, "plz"), _k("ort", kcols, "ort"),
            _k("email", kcols, "email"), _k("telefon", kcols, "telefon"),
            _k("kundentyp", kcols, "pk_zw"), _k("rechnungsart", kcols, "rechnungsart"),
        ])

    def _q_verkaeufe(self, con, von, bis, kcols):
        sql = f"""
            SELECT b.kundennummer AS kundennummer,
                   COALESCE(k.kundenname, b.apotheke) AS apotheke,
                   {self._kunden_select(kcols)},
                   p.pzn AS pzn, p.artikelname AS artikel,
                   SUM(p.menge) AS menge,
                   SUM({_NETTO}) AS umsatz_artikel
            FROM tbl_bestellpositionen p
            JOIN tbl_bestellungen b ON b.id = p.bestell_id
            LEFT JOIN tbl_kunden_center k ON k.kundennummer = b.kundennummer
            WHERE date(b.datum) BETWEEN date(?) AND date(?)
              AND COALESCE(b.status,'') NOT IN ('storniert', 'abgesagt')
              AND COALESCE(p.bestellart,'Bestellung') = 'Bestellung'
            GROUP BY b.kundennummer, p.pzn, p.artikelname
            ORDER BY apotheke, artikel
        """
        rows = [dict(r) for r in con.execute(sql, (von, bis))]
        # Umsatz pro Kunde (Summe ueber alle Artikel) als Zusatzspalte
        totals: dict[str, float] = {}
        for r in rows:
            totals[r["kundennummer"]] = totals.get(r["kundennummer"], 0.0) + (r["umsatz_artikel"] or 0.0)
        for r in rows:
            r["umsatz_kunde"] = totals.get(r["kundennummer"], 0.0)
        return rows

    def _q_kunden(self, con, von, bis, kcols):
        sql = f"""
            SELECT b.kundennummer AS kundennummer,
                   COALESCE(k.kundenname, b.apotheke) AS apotheke,
                   {self._kunden_select(kcols)},
                   COUNT(DISTINCT p.pzn) AS anzahl_artikel,
                   SUM(p.menge) AS gesamt_menge,
                   SUM({_NETTO}) AS gesamt_umsatz
            FROM tbl_bestellpositionen p
            JOIN tbl_bestellungen b ON b.id = p.bestell_id
            LEFT JOIN tbl_kunden_center k ON k.kundennummer = b.kundennummer
            WHERE date(b.datum) BETWEEN date(?) AND date(?)
              AND COALESCE(b.status,'') NOT IN ('storniert', 'abgesagt')
              AND COALESCE(p.bestellart,'Bestellung') = 'Bestellung'
            GROUP BY b.kundennummer
            ORDER BY gesamt_umsatz DESC
        """
        return [dict(r) for r in con.execute(sql, (von, bis))]

    def _q_artikel(self, con, von, bis):
        sql = f"""
            SELECT p.pzn AS pzn, p.artikelname AS artikel,
                   COUNT(DISTINCT b.kundennummer) AS anzahl_kunden,
                   SUM(p.menge) AS gesamt_menge,
                   SUM({_NETTO}) AS gesamt_umsatz
            FROM tbl_bestellpositionen p
            JOIN tbl_bestellungen b ON b.id = p.bestell_id
            WHERE date(b.datum) BETWEEN date(?) AND date(?)
              AND COALESCE(b.status,'') NOT IN ('storniert', 'abgesagt')
              AND COALESCE(p.bestellart,'Bestellung') = 'Bestellung'
            GROUP BY p.pzn, p.artikelname
            ORDER BY gesamt_umsatz DESC
        """
        return [dict(r) for r in con.execute(sql, (von, bis))]

    def _q_frei(self, con, von, bis, kcols):
        """Freie Auswertung: SQL dynamisch aus Dimensionen + Kennzahlen + Filtern.
        Alle Ausdruecke stammen aus den Whitelists DIM_MAP/MEA_MAP/FILTER_MAP;
        Werte werden gebunden. Rueckgabe: (rows, result_cols)."""
        dim_keys = [k for k in self.builder.get("dims", []) if k in DIM_MAP]
        mea_keys = [k for k in self.builder.get("measures", []) if k in MEA_MAP] or ["umsatz", "menge"]
        select, group, result_cols = [], [], []
        for k in dim_keys:
            _, label, sql, kcol = DIM_MAP[k]
            expr = sql if (kcol is None or kcol in kcols) else "NULL"
            select.append(f"{expr} AS {k}")
            group.append(expr)
            result_cols.append((k, label, "text"))
        for k in mea_keys:
            _, label, sql, typ = MEA_MAP[k]
            select.append(f"{sql} AS {k}")
            result_cols.append((k, label, typ))

        where = ["date(b.datum) BETWEEN date(?) AND date(?)",
                 "COALESCE(b.status,'') NOT IN ('storniert', 'abgesagt')",
                 "COALESCE(p.bestellart,'Bestellung') = 'Bestellung'"]
        params = [von, bis]
        having, hparams = [], []
        for spec in self.builder.get("filters", []):
            built = _build_filter(spec, kcols)
            if not built:
                continue
            clause, ps, scope = built
            if scope == "having":
                having.append(clause); hparams += ps
            else:
                where.append(clause); params += ps

        sql = (f"SELECT {', '.join(select) or '1'} "
               "FROM tbl_bestellpositionen p "
               "JOIN tbl_bestellungen b ON b.id = p.bestell_id "
               "LEFT JOIN tbl_kunden_center k ON k.kundennummer = b.kundennummer "
               f"WHERE {' AND '.join(where)}")
        if group:
            sql += " GROUP BY " + ", ".join(group)
        if having:
            sql += " HAVING " + " AND ".join(having)
        sql += f" ORDER BY {MEA_MAP[mea_keys[0]][2]} DESC"
        rows = [dict(r) for r in con.execute(sql, params + hparams)]
        return rows, result_cols

    # ── Anzeige ──────────────────────────────────────────────────────────────
    def _active_cols(self):
        """(key, label, typ) der aktuell gewaehlten Spalten in Reihenfolge."""
        if self.perspektive == "frei":
            return list(self._result_cols)
        defs = {k: (lab, typ) for k, lab, typ in COLS[self.perspektive]}
        out = []
        for key in self.selected[self.perspektive]:
            if key in defs:
                out.append((key, defs[key][0], defs[key][1]))
        return out

    def _fmt(self, value, typ):
        if typ == "eur":
            return _eur(value)
        if typ == "int":
            try:
                return str(int(value or 0))
            except (TypeError, ValueError):
                return str(value or "")
        if typ == "num":
            try:
                return f"{float(value):.1f}"
            except (TypeError, ValueError):
                return "" if value is None else str(value)
        return "" if value is None else str(value)

    def _render(self, rows, status_text):
        cols = self._active_cols()
        keys = [c[0] for c in cols]
        self.tree.delete(*self.tree.get_children())
        self.tree["columns"] = keys
        for key, label, typ in cols:
            self.tree.heading(key, text=label, command=lambda k=key: self._sort_by(k))
            self.tree.column(key, width=140 if typ != "text" else 168, minwidth=70,
                             anchor="e" if typ in ("eur", "int", "num") else "w")
        for i, r in enumerate(rows):
            self.tree.insert("", "end", tags=("odd" if i % 2 else "even",),
                             values=[self._fmt(r.get(k), t) for k, _, t in cols])
        self.status.configure(text=status_text)

    def _sort_by(self, key):
        typ = dict((k, t) for k, _, t in self._active_cols()).get(key, "text")

        def sort_key(r):
            v = r.get(key)
            if typ in ("eur", "int"):
                try:
                    return (0, float(v or 0))
                except (TypeError, ValueError):
                    return (0, 0.0)
            return (1, str(v or "").lower())

        reverse = getattr(self, "_sort_state", {}).get(key, False)
        self._rows.sort(key=sort_key, reverse=not reverse)
        self._sort_state = {key: not reverse}
        self._render(self._rows, self.status.cget("text"))

    # ── Spalten-Dialog (Dual-Liste) ──────────────────────────────────────────
    def _open_columns_dialog(self):
        persp = self.perspektive
        all_defs = [(k, lab) for k, lab, _ in COLS[persp]]
        chosen = list(self.selected[persp])

        win = tk.Toplevel(self)
        win.title("Ausgabespalten wählen")
        win.configure(bg=SHELL_BG)
        win.transient(self.winfo_toplevel())
        win.grab_set()

        tk.Label(win, text="Ausgabespalten", bg=SHELL_BG, fg=ACCENT,
                 font=("Segoe UI", 15, "bold")).grid(row=0, column=0, sticky="w", padx=18, pady=(16, 0))
        tk.Label(win, text="Per Drag & Drop hinzufügen, sortieren oder zurückziehen — oder die Pfeile / Doppelklick nutzen.",
                 bg=SHELL_BG, fg=MUTED, font=FONT_BASE).grid(row=1, column=0, sticky="w", padx=18, pady=(3, 12))
        body = tk.Frame(win, bg=SHELL_BG)
        body.grid(row=2, column=0, padx=18)
        refresh = self._make_dual_list(body, all_defs, chosen)

        def apply_and_close():
            if not chosen:
                messagebox.showwarning("Spalten", "Mindestens eine Spalte wählen.", parent=win)
                return
            self.selected[persp] = list(chosen)
            win.destroy()
            self._render(self._rows, self.status.cget("text"))

        bar = tk.Frame(win, bg=SHELL_BG)
        bar.grid(row=3, column=0, pady=(14, 16))
        self._flat_btn(bar, "Übernehmen", apply_and_close, bg=ACCENT, big=True).pack(side="left", padx=6)
        self._flat_btn(bar, "Standard", lambda: (chosen.clear(), chosen.extend(DEFAULTS[persp]), refresh()),
                       bg="#e8eef6", fg=ACCENT).pack(side="left", padx=6)
        self._flat_btn(bar, "Abbrechen", win.destroy, bg="#e8eef6", fg=ACCENT).pack(side="left", padx=6)

    # ── Freie Auswertung: Bauen-Dialog + Vorlagen ────────────────────────────
    def _vorlagen_dir(self) -> Path:
        d = OUTPUT_DIR / "Auswertungen" / "vorlagen"
        try:
            d.mkdir(parents=True, exist_ok=True)
        except Exception:
            pass
        return d

    def _list_vorlagen(self):
        try:
            return sorted(p.stem for p in self._vorlagen_dir().glob("*.json"))
        except Exception:
            return []

    def _make_dual_list(self, parent, all_defs, chosen, headers=("Verfügbar", "Ausgewählt")):
        """Zwei Listen mit DRAG & DROP. all_defs=[(key,label)]; 'chosen' (keys) wird
        in-place gepflegt. Gibt refresh() zurueck.
          - aus 'Verfügbar' nach 'Ausgewählt' ziehen  = hinzufuegen (an Drop-Position)
          - innerhalb 'Ausgewählt' ziehen             = sortieren (live)
          - aus 'Ausgewählt' zurueck nach 'Verfügbar' = entfernen
        Zusaetzlich Doppelklick und →/←/↑/↓-Buttons.
        """
        bg = parent["bg"]
        label_of = {k: lab for k, lab in all_defs}
        lbopt = dict(exportselection=False, font=FONT_LIST, activestyle="none",
                     bd=1, relief="solid", highlightthickness=1,
                     highlightbackground=CARD_BORDER, selectbackground=NAV_SEL,
                     selectforeground=ACCENT, height=12, width=24)
        tk.Label(parent, text=headers[0], bg=bg, fg=MUTED,
                 font=("Segoe UI", 9, "bold")).grid(row=0, column=0, sticky="w", pady=(0, 2))
        tk.Label(parent, text=headers[1], bg=bg, fg=MUTED,
                 font=("Segoe UI", 9, "bold")).grid(row=0, column=2, sticky="w", pady=(0, 2))
        lb_av = tk.Listbox(parent, **lbopt)
        lb_av.grid(row=1, column=0, rowspan=5, padx=(0, 6), pady=2)
        lb_se = tk.Listbox(parent, **lbopt)
        lb_se.grid(row=1, column=2, rowspan=5, padx=(6, 0), pady=2)

        def avail():
            return [k for k, _ in all_defs if k not in chosen]

        def refresh():
            lb_av.delete(0, tk.END); lb_se.delete(0, tk.END)
            for k in avail():
                lb_av.insert(tk.END, label_of[k])
            for k in chosen:
                lb_se.insert(tk.END, label_of.get(k, k))

        def add_key(key, at=None):
            if key in chosen:
                return
            if at is None or at >= len(chosen):
                chosen.append(key)
            else:
                chosen.insert(max(0, at), key)
            refresh()

        def add():
            s = lb_av.curselection()
            if s:
                add_key(avail()[s[0]])

        def rem():
            s = lb_se.curselection()
            if s:
                chosen.pop(s[0]); refresh()

        def mv(d):
            s = lb_se.curselection()
            if not s:
                return
            i = s[0]; j = i + d
            if 0 <= j < len(chosen):
                chosen[i], chosen[j] = chosen[j], chosen[i]; refresh(); lb_se.selection_set(j)

        # ── Drag & Drop ──
        drag = {"key": None, "from": None}

        def on_press(ev):
            w = ev.widget
            idx = w.nearest(ev.y)
            keys = avail() if w is lb_av else chosen
            if 0 <= idx < len(keys):
                drag["key"] = keys[idx]; drag["from"] = w
                w.selection_clear(0, tk.END); w.selection_set(idx)

        def on_motion(ev):
            # Live-Sortierung innerhalb 'Ausgewählt'
            if drag["from"] is lb_se and drag["key"] in chosen:
                j = lb_se.nearest(ev.y)
                i = chosen.index(drag["key"])
                if 0 <= j < len(chosen) and j != i:
                    chosen.insert(j, chosen.pop(i)); refresh(); lb_se.selection_set(j)

        def on_release(ev):
            key, frm = drag["key"], drag["from"]
            drag["key"] = drag["from"] = None
            if not key:
                return
            tgt = parent.winfo_containing(ev.x_root, ev.y_root)
            if frm is lb_av and tgt is lb_se:
                at = lb_se.nearest(ev.y_root - lb_se.winfo_rooty())
                add_key(key, at if at >= 0 else None)
            elif frm is lb_se and tgt is lb_av and key in chosen:
                chosen.remove(key); refresh()

        for lb in (lb_av, lb_se):
            lb.bind("<ButtonPress-1>", on_press)
            lb.bind("<B1-Motion>", on_motion)
            lb.bind("<ButtonRelease-1>", on_release)
        lb_av.bind("<Double-Button-1>", lambda e: add())
        lb_se.bind("<Double-Button-1>", lambda e: rem())

        bf = tk.Frame(parent, bg=bg)
        bf.grid(row=1, column=1, rowspan=5, padx=4)
        for t, c in [("→", add), ("←", rem), ("↑", lambda: mv(-1)), ("↓", lambda: mv(1))]:
            tk.Button(bf, text=t, width=3, relief="flat", bg="#e8eef6", fg=ACCENT,
                      cursor="hand2", command=c).pack(pady=3)
        refresh()
        return refresh

    def _open_builder_dialog(self):
        dims = list(self.builder.get("dims", []))
        meas = list(self.builder.get("measures", []))

        win = tk.Toplevel(self)
        win.title("Auswertung bauen")
        win.configure(bg=SHELL_BG)
        win.transient(self.winfo_toplevel()); win.grab_set()

        tk.Label(win, text="Gruppieren nach (Dimensionen)", bg=SHELL_BG, fg=ACCENT,
                 font=("Segoe UI", 10, "bold")).grid(row=0, column=0, sticky="w", padx=10, pady=(10, 2))
        dframe = tk.Frame(win, bg=SHELL_BG)
        dframe.grid(row=1, column=0, padx=10, sticky="nw")
        self._make_dual_list(dframe, [(k, lab) for k, lab, *_ in DIMENSIONS], dims)

        tk.Label(win, text="Kennzahlen (Werte)", bg=SHELL_BG, fg=ACCENT,
                 font=("Segoe UI", 10, "bold")).grid(row=0, column=1, sticky="w", padx=10, pady=(10, 2))
        mframe = tk.Frame(win, bg=SHELL_BG)
        mframe.grid(row=1, column=1, padx=10, sticky="nw")
        self._make_dual_list(mframe, [(k, lab) for k, lab, *_ in MEASURES], meas)

        tk.Label(win, text="Filter (alle müssen zutreffen)", bg=SHELL_BG, fg=ACCENT,
                 font=("Segoe UI", 10, "bold")).grid(row=2, column=0, columnspan=2, sticky="w", padx=10, pady=(10, 2))
        fwrap = tk.Frame(win, bg=SHELL_BG)
        fwrap.grid(row=3, column=0, columnspan=2, sticky="w", padx=10)
        rows_widgets = []
        flabels = [lab for _, lab, *_ in FILTER_FIELDS]
        fkey_by_label = {lab: k for k, lab, *_ in FILTER_FIELDS}
        fkind_by_label = {lab: kind for _, lab, _, kind, _, _ in FILTER_FIELDS}

        def add_filter_row(init=None):
            rowf = tk.Frame(fwrap, bg=SHELL_BG)
            rowf.grid(sticky="w", pady=2)
            fld = ttk.Combobox(rowf, values=flabels, width=20, state="readonly")
            opv = ttk.Combobox(rowf, values=OPS_TEXT, width=11, state="readonly")
            valv = tk.Entry(rowf, width=16)

            def on_field(*_):
                kind = fkind_by_label.get(fld.get(), "text")
                ops = OPS_NUM if kind == "num" else OPS_TEXT
                opv.configure(values=ops)
                if opv.get() not in ops:
                    opv.set(ops[0])

            fld.bind("<<ComboboxSelected>>", on_field)
            fld.pack(side="left", padx=2); opv.pack(side="left", padx=2); valv.pack(side="left", padx=2)
            entry = {"frame": rowf, "fld": fld, "op": opv, "val": valv}

            def remove():
                rowf.destroy()
                if entry in rows_widgets:
                    rows_widgets.remove(entry)

            tk.Button(rowf, text="✕", width=2, command=remove).pack(side="left", padx=2)
            rows_widgets.append(entry)
            if init:
                lab = next((l for k, l, *_ in FILTER_FIELDS if k == init.get("field")), "")
                fld.set(lab); on_field()
                opv.set(init.get("op") or opv.get()); valv.insert(0, str(init.get("value", "")))

        for f in self.builder.get("filters", []):
            add_filter_row(f)
        tk.Button(win, text="+ Filter", command=lambda: add_filter_row()).grid(
            row=4, column=0, sticky="w", padx=10, pady=(2, 8))

        def collect():
            specs = []
            for e in rows_widgets:
                lab = e["fld"].get()
                if not lab:
                    continue
                specs.append({"field": fkey_by_label.get(lab), "op": e["op"].get(), "value": e["val"].get()})
            return {"dims": list(dims), "measures": list(meas), "filters": specs}

        # Vorlagen
        vrow = tk.Frame(win, bg=SHELL_BG)
        vrow.grid(row=5, column=0, columnspan=2, sticky="w", padx=10, pady=(0, 6))
        tk.Label(vrow, text="Vorlage:", bg=SHELL_BG).pack(side="left")
        vcombo = ttk.Combobox(vrow, values=self._list_vorlagen(), width=22, state="readonly")
        vcombo.pack(side="left", padx=4)

        def do_save():
            name = simpledialog.askstring("Vorlage speichern", "Name der Vorlage:", parent=win)
            if not name:
                return
            safe = "".join(c if c.isalnum() or c in " -_" else "_" for c in name).strip() or "Vorlage"
            try:
                (self._vorlagen_dir() / f"{safe}.json").write_text(
                    json.dumps(collect(), ensure_ascii=False, indent=2), encoding="utf-8")
                vcombo.configure(values=self._list_vorlagen()); vcombo.set(safe)
                messagebox.showinfo("Vorlage", _T('Gespeichert: {p0}', p0=safe), parent=win)
            except Exception as exc:
                messagebox.showerror("Vorlage", _T('Speichern fehlgeschlagen:\n{p0}', p0=exc), parent=win)

        def do_load():
            name = vcombo.get()
            if not name:
                return
            try:
                data = json.loads((self._vorlagen_dir() / f"{name}.json").read_text(encoding="utf-8"))
            except Exception as exc:
                messagebox.showerror("Vorlage", _T('Laden fehlgeschlagen:\n{p0}', p0=exc), parent=win)
                return
            self.builder = {
                "dims": [k for k in data.get("dims", []) if k in DIM_MAP],
                "measures": [k for k in data.get("measures", []) if k in MEA_MAP],
                "filters": [f for f in data.get("filters", []) if f.get("field") in FILTER_MAP],
            }
            win.destroy(); self._open_builder_dialog()

        tk.Button(vrow, text="Laden", command=do_load).pack(side="left", padx=2)
        tk.Button(vrow, text="Speichern als…", command=do_save).pack(side="left", padx=2)

        bar = tk.Frame(win, bg=SHELL_BG)
        bar.grid(row=6, column=0, columnspan=2, pady=(4, 10))

        def apply_close():
            self.builder = collect()
            win.destroy(); self.run_query()

        tk.Button(bar, text="Übernehmen & Auswerten", bg=ACCENT, fg="white", relief="flat",
                  padx=14, pady=4, command=apply_close).pack(side="left", padx=6)
        tk.Button(bar, text="Abbrechen", relief="flat", padx=10, pady=4,
                  command=win.destroy).pack(side="left", padx=6)

    # ── Export ───────────────────────────────────────────────────────────────
    def _export_matrix(self):
        cols = self._active_cols()
        headers = [c[1] for c in cols]
        # Rohwerte fuer Excel (Zahlen bleiben Zahlen), formatiert fuer Word/PDF
        raw, disp = [], []
        for r in self._rows:
            raw.append([r.get(k) for k, _, _ in cols])
            disp.append([self._fmt(r.get(k), t) for k, _, t in cols])
        return headers, raw, disp, [t for _, _, t in cols]

    def export(self, fmt):
        if not self._rows:
            messagebox.showinfo("Export", "Keine Daten zum Exportieren. Erst auswerten.", parent=self)
            return
        headers, raw, disp, types = self._export_matrix()
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        label = dict(PERSPEKTIVEN)[self.perspektive]
        try:
            OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
            base = OUTPUT_DIR / "Auswertungen"
            base.mkdir(parents=True, exist_ok=True)
        except Exception:
            base = Path.home()
        try:
            if fmt == "excel":
                out = self._export_excel(base / f"Auswertung_{label}_{stamp}.xlsx", headers, raw, types)
            elif fmt == "word":
                out = self._export_word(base / f"Auswertung_{label}_{stamp}.docx", headers, disp, label)
            else:
                out = self._export_pdf(base / f"Auswertung_{label}_{stamp}.pdf", headers, disp, label)
        except ImportError as exc:
            messagebox.showwarning(
                "Bibliothek fehlt",
                _T('Für diesen Export fehlt eine Bibliothek:\n{p0}\n\nInstallieren mit:\n  pip install python-docx   (Word)\n  pip install reportlab      (PDF)', p0=exc), parent=self)
            return
        except Exception as exc:
            messagebox.showerror("Export", _T('Export fehlgeschlagen:\n{p0}', p0=exc), parent=self)
            return
        if messagebox.askyesno("Export fertig", _T('Gespeichert:\n{p0}\n\nJetzt öffnen?', p0=out), parent=self):
            try:
                os.startfile(out)  # type: ignore[attr-defined]
            except Exception:
                pass

    def _export_excel(self, path, headers, raw, types):
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
        wb = Workbook()
        ws = wb.active
        ws.title = "Auswertung"
        ws.append(headers)
        for c in ws[1]:
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = PatternFill("solid", fgColor="0B4A86")
        for row in raw:
            ws.append([("" if v is None else v) for v in row])
        # Euro-Format auf eur-Spalten
        for ci, typ in enumerate(types, start=1):
            if typ == "eur":
                for r in range(2, ws.max_row + 1):
                    ws.cell(row=r, column=ci).number_format = '#,##0.00 €'
        for ci, h in enumerate(headers, start=1):
            ws.column_dimensions[ws.cell(row=1, column=ci).column_letter].width = max(12, len(str(h)) + 2)
        wb.save(path)
        return path

    def _export_word(self, path, headers, disp, title):
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_heading(f"Auswertung · {title}", level=1)
        doc.add_paragraph(datetime.now().strftime("Erstellt am %d.%m.%Y %H:%M"))
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = str(h)
        for row in disp:
            cells = table.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = "" if v is None else str(v)
        doc.save(path)
        return path

    def _export_pdf(self, path, headers, disp, title):
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        styles = getSampleStyleSheet()
        doc = SimpleDocTemplate(str(path), pagesize=landscape(A4))
        data = [headers] + [["" if v is None else str(v) for v in row] for row in disp]
        t = Table(data, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B4A86")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#eef3fa")]),
        ]))
        doc.build([Paragraph(f"Auswertung · {title}", styles["Title"]), Spacer(1, 8), t])
        return path

    # ── NMGone-Verbindung ────────────────────────────────────────────────────
    def _open_nmgone(self):
        if self._nmgone_action:
            try:
                self._nmgone_action()
                return
            except Exception:
                pass
        try:
            import subprocess
            if getattr(sys, "frozen", False):
                subprocess.Popen([sys.executable])
            else:
                start_py = Path(__file__).resolve().parent.parent / "start.py"
                subprocess.Popen([sys.executable, str(start_py)])
        except Exception as exc:
            messagebox.showerror("NMGone", _T('NMGone konnte nicht gestartet werden:\n{p0}', p0=exc), parent=self)


def run_standalone():
    """Startet die Auswertungen als eigenes Fenster (eigenes Taskleisten-Icon).
    Genutzt von start_report.py und von NMGone.exe --report."""
    if os.name == "nt":
        try:
            import ctypes
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("NMG.Report")
        except Exception:
            pass
    try:
        from .db import init_db
        init_db(DB_PATH)
    except Exception:
        pass
    root = tk.Tk()
    root.title(f"NMG Auswertungen{DEMO_SUFFIX}")
    root.geometry("1280x820")
    root.minsize(1040, 680)
    # Im Vollbild (maximiert) starten. 'zoomed' = Windows; sonst -zoomed/Bildschirmgroesse.
    try:
        root.state("zoomed")
    except tk.TclError:
        try:
            root.attributes("-zoomed", True)
        except tk.TclError:
            root.geometry(f"{root.winfo_screenwidth()}x{root.winfo_screenheight()}+0+0")
    root.configure(bg=SHELL_BG)
    theme.apply_theme(root)
    theme.apply_widget_defaults(root)
    for ico in ("report.ico", "NMGone.ico"):
        try:
            root.iconbitmap(str(ASSETS_DIR / ico))
            break
        except Exception:
            continue
    BerichtPanel(root, on_close=root.destroy).pack(fill="both", expand=True)
    root.mainloop()


if __name__ == "__main__":
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
    run_standalone()
