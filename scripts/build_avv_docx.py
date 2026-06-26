# -*- coding: utf-8 -*-
"""Erzeugt ein unterschriftsfertiges Word-Dokument:

  Auftragsverarbeitungsvertrag (AVV) nach Art. 28 DSGVO
  + Anlage 1: Technische und organisatorische Massnahmen (Art. 32 DSGVO)
  + Anlage 2: Genehmigte Unterauftragsverarbeiter

Kontext: Anbieter (= unsere Firma, Inhaber unterschreibt) hostet die
Personal-/Mitarbeiterdaten der Kundenfirma im Web. Der Kunde ist
Verantwortlicher, wir sind Auftragsverarbeiter.

WICHTIG: Mustervorlage, KEINE Rechtsberatung. Vor produktivem Einsatz von
einem Anwalt / Datenschutzbeauftragten pruefen lassen.

Aufruf:  python scripts/build_avv_docx.py
Bedarf:  pip install python-docx
"""
from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "docs", "AV-Vertrag_Auftragsverarbeitung.docx")

BLAU = RGBColor(0x0B, 0x4A, 0x86)
GRAU = RGBColor(0x5A, 0x66, 0x75)
ROT = RGBColor(0xB5, 0x39, 0x1F)

doc = Document()

# --- Grundlayout: A4, Arial, Raender ----------------------------------------
sec = doc.sections[0]
sec.page_height = Mm(297)
sec.page_width = Mm(210)
sec.top_margin = Mm(22)
sec.bottom_margin = Mm(20)
sec.left_margin = Mm(22)
sec.right_margin = Mm(22)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12
# Ostasiatische Schrift ebenfalls auf Arial (saubere Darstellung)
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
rfonts.set(qn("w:eastAsia"), "Arial")


def set_cell_bg(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcpr.append(shd)


def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = BLAU
    p.paragraph_format.space_after = Pt(2)
    return p


def subtitle(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = GRAU
    p.paragraph_format.space_after = Pt(10)
    return p


def h(text, color=BLAU, size=12, space_before=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(text, size=10, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(items, size=10):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(it)
        r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(3)


def note_box(lines, fill="FCF3E6"):
    """Einspaltige Box (Tabelle mit 1 Zelle) fuer Hinweise."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_cell_bg(cell, fill)
    cell.paragraphs[0].text = ""
    first = True
    for ln in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        bold = ln.get("bold", False)
        r = p.add_run(ln["t"])
        r.bold = bold
        r.font.size = Pt(ln.get("size", 9))
        if ln.get("color"):
            r.font.color.rgb = ln["color"]
        p.paragraph_format.space_after = Pt(2)
    _set_table_borders(t, "D7DEE6")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _set_table_borders(table, color="D7DEE6", size="6"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def kv_table(rows, widths=(55, 123), head=None, headfill="0B4A86"):
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if head:
        hr = t.add_row().cells
        for i, txt in enumerate(head):
            set_cell_bg(hr[i], headfill)
            rp = hr[i].paragraphs[0].add_run(txt)
            rp.bold = True
            rp.font.size = Pt(9)
            rp.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for k, v in rows:
        cells = t.add_row().cells
        rk = cells[0].paragraphs[0].add_run(k)
        rk.bold = True
        rk.font.size = Pt(9)
        rv = cells[1].paragraphs[0].add_run(v)
        rv.font.size = Pt(9)
    for i, w in enumerate(widths):
        for row in t.rows:
            row.cells[i].width = Mm(w)
    _set_table_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def fill(label):
    """Ausfuell-Feld mit Platzhalter-Linie."""
    return f"{label}: ____________________________________________"


# ============================================================ DECKBLATT
title("Vertrag zur Auftragsverarbeitung (AVV)")
subtitle("gemäß Art. 28 der Datenschutz-Grundverordnung (DSGVO)")

note_box([
    {"t": "Mustervorlage – keine Rechtsberatung", "bold": True, "size": 10, "color": ROT},
    {"t": "Dieses Dokument ist eine sorgfältig erstellte, marktübliche Vorlage. Es ersetzt "
          "keine anwaltliche Prüfung. Vor dem produktiven Einsatz mit echten Kundendaten "
          "von einem Rechtsanwalt oder Datenschutzbeauftragten prüfen und an den konkreten "
          "Einzelfall anpassen lassen.", "size": 9},
], fill="FCEDEA")

h("Zwischen den Vertragsparteien", size=12, space_before=10)

body("Verantwortlicher (Auftraggeber / Kunde) – nachfolgend „Verantwortlicher“:",
     bold=True, space_after=3)
kv_table([
    (fill("Firma / Name").split(":")[0], "____________________________________________"),
    ("Anschrift", "____________________________________________"),
    ("Vertreten durch", "____________________________________________"),
    ("E-Mail / Telefon", "____________________________________________"),
], head=None)

body("Auftragsverarbeiter (Anbieter) – nachfolgend „Auftragsverarbeiter“:",
     bold=True, space_after=3)
kv_table([
    ("Firma / Name", "[Anbieter-Firma], Inhaber: [Vor- und Nachname]"),
    ("Anschrift", "[Straße, PLZ, Ort]"),
    ("Vertreten durch", "[Inhaber / Geschäftsführung]"),
    ("E-Mail / Telefon", "[Kontakt]"),
], head=None)

body("– nachfolgend gemeinsam die „Parteien“ –", italic=True, color=GRAU,
     space_after=8)

# ---- Praeambel
h("Präambel")
body("Der Auftragsverarbeiter stellt dem Verantwortlichen eine webbasierte Software zur "
     "Verwaltung von Mitarbeiter- und Personaldaten („Personal-Modul“) als "
     "Software-as-a-Service bereit und betreibt die hierfür erforderliche IT-Infrastruktur. "
     "Dabei verarbeitet der Auftragsverarbeiter personenbezogene Daten ausschließlich im "
     "Auftrag und nach Weisung des Verantwortlichen. Dieser Vertrag konkretisiert die "
     "datenschutzrechtlichen Pflichten der Parteien gemäß Art. 28 DSGVO. Er ist Bestandteil "
     "des zwischen den Parteien geschlossenen Hauptvertrags (Nutzungs-/Dienstleistungsvertrag).")

# ============================================================ PARAGRAFEN
h("§ 1 Gegenstand und Dauer des Auftrags")
numbered([
    "Gegenstand des Auftrags ist die Verarbeitung personenbezogener Daten durch den "
    "Auftragsverarbeiter für den Verantwortlichen im Rahmen der Bereitstellung und des "
    "Betriebs des Personal-Moduls (Hosting, Speicherung, Wartung, Support).",
    "Die Dauer dieses Vertrags entspricht der Laufzeit des Hauptvertrags. Er endet automatisch "
    "mit dessen Beendigung, ohne dass es einer gesonderten Kündigung bedarf; die Pflichten "
    "zur Löschung/Rückgabe nach § 10 bleiben hiervon unberührt.",
])

h("§ 2 Art, Umfang und Zweck der Verarbeitung; betroffene Personen und Datenarten")
body("Der Auftragsverarbeiter verarbeitet personenbezogene Daten im folgenden Rahmen. "
     "Eine Übermittlung in ein Drittland außerhalb der EU/des EWR findet nicht statt.",
     space_after=4)
kv_table([
    ("Zweck der Verarbeitung", "Bereitstellung, Betrieb und Wartung der Personalverwaltungs-Software "
                               "im Auftrag des Verantwortlichen"),
    ("Art der Verarbeitung", "Erheben, Speichern, Anzeigen, Ändern, Löschen, Sichern (Backup), "
                             "Übermitteln innerhalb der Anwendung"),
    ("Kategorien betroffener Personen", "Beschäftigte des Verantwortlichen (Mitarbeitende, ggf. "
                                        "Bewerber), Vorgesetzte"),
    ("Kategorien personenbezogener Daten", "Stammdaten (Name, Kontakt), Personal-/Organisationsdaten "
                                           "(Funktion, Vorgesetzte, Eintritt), frei definierbare "
                                           "Zusatzfelder; ggf. weitere vom Verantwortlichen eingegebene Daten"),
    ("Besondere Kategorien (Art. 9)", "Nur, soweit der Verantwortliche solche Daten eigenverantwortlich "
                                      "eingibt (z. B. Gesundheitsdaten); dann mit erhöhten Schutzmaßnahmen"),
], head=["Merkmal", "Beschreibung"])

h("§ 3 Weisungsrecht des Verantwortlichen")
numbered([
    "Der Auftragsverarbeiter verarbeitet die Daten ausschließlich im Rahmen der getroffenen "
    "Vereinbarungen und nach dokumentierten Weisungen des Verantwortlichen, es sei denn, er ist "
    "gesetzlich zu einer anderweitigen Verarbeitung verpflichtet (Art. 28 Abs. 3 lit. a DSGVO).",
    "Weisungen erfolgen grundsätzlich in Textform. Mündliche Weisungen sind unverzüglich "
    "in Textform zu bestätigen.",
    "Der Auftragsverarbeiter informiert den Verantwortlichen unverzüglich, wenn er der Auffassung "
    "ist, dass eine Weisung gegen geltendes Datenschutzrecht verstößt (Art. 28 Abs. 3 S. 3 DSGVO). "
    "Er ist berechtigt, die Durchführung der betreffenden Weisung auszusetzen, bis sie bestätigt "
    "oder geändert wird.",
])

h("§ 4 Pflichten des Auftragsverarbeiters")
numbered([
    "Vertraulichkeit: Der Auftragsverarbeiter setzt zur Verarbeitung nur Beschäftigte ein, die auf "
    "das Datengeheimnis verpflichtet und mit den einschlägigen Datenschutzbestimmungen vertraut "
    "gemacht wurden (Art. 28 Abs. 3 lit. b, Art. 29, Art. 32 Abs. 4 DSGVO).",
    "Datensicherheit: Der Auftragsverarbeiter trifft die technischen und organisatorischen Maßnahmen "
    "nach Art. 32 DSGVO gemäß Anlage 1 und hält diese während der Vertragslaufzeit auf dem "
    "Stand der Technik.",
    "Unterstützung: Der Auftragsverarbeiter unterstützt den Verantwortlichen im Rahmen seiner "
    "Möglichkeiten bei der Erfüllung der Betroffenenrechte (§ 7) sowie der Pflichten nach "
    "Art. 32 bis 36 DSGVO (Sicherheit, Meldepflichten, Datenschutz-Folgenabschätzung).",
    "Meldung: Der Auftragsverarbeiter meldet Verletzungen des Schutzes personenbezogener Daten gemäß "
    "§ 8.",
    "Datenschutzbeauftragter: Sofern gesetzlich erforderlich, benennt der Auftragsverarbeiter einen "
    "Datenschutzbeauftragten und teilt dessen Kontaktdaten mit. Andernfalls benennt er eine für den "
    "Datenschutz zuständige Ansprechperson.",
    "Nachweis: Der Auftragsverarbeiter stellt dem Verantwortlichen alle erforderlichen Informationen "
    "zum Nachweis der Einhaltung der Pflichten aus Art. 28 DSGVO zur Verfügung und ermöglicht "
    "Überprüfungen gemäß § 9.",
    "Verzeichnis: Der Auftragsverarbeiter führt ein Verzeichnis aller Kategorien von im Auftrag "
    "durchgeführten Verarbeitungstätigkeiten nach Art. 30 Abs. 2 DSGVO.",
])

h("§ 5 Technische und organisatorische Maßnahmen (TOM)")
numbered([
    "Der Auftragsverarbeiter gewährleistet die in Anlage 1 beschriebenen technischen und "
    "organisatorischen Maßnahmen nach Art. 32 DSGVO.",
    "Die Maßnahmen können im Lauf des Vertragsverhältnisses an den technischen Fortschritt "
    "angepasst werden, dürfen das vereinbarte Schutzniveau jedoch nicht unterschreiten. Wesentliche "
    "Änderungen werden dem Verantwortlichen mitgeteilt.",
])

h("§ 6 Unterauftragsverhältnisse")
numbered([
    "Der Verantwortliche stimmt dem Einsatz der in Anlage 2 genannten Unterauftragsverarbeiter zu "
    "(allgemeine schriftliche Genehmigung gemäß Art. 28 Abs. 2 DSGVO).",
    "Ein Wechsel oder die Hinzunahme weiterer Unterauftragsverarbeiter wird dem Verantwortlichen "
    "vorab in Textform mitgeteilt. Der Verantwortliche kann einer Änderung aus wichtigem, "
    "datenschutzrechtlichem Grund innerhalb von 14 Tagen widersprechen.",
    "Der Auftragsverarbeiter verpflichtet jeden Unterauftragsverarbeiter vertraglich auf "
    "Datenschutzpflichten, die denen dieses Vertrags entsprechen (Art. 28 Abs. 4 DSGVO).",
    "Nicht als Unterauftragsverarbeitung gelten Neben-/Hilfsleistungen Dritter (z. B. Telekommunikation, "
    "Wartung). Insoweit stellt der Auftragsverarbeiter angemessene vertragliche Vereinbarungen und "
    "Kontrollmaßnahmen sicher.",
])

h("§ 7 Wahrung der Betroffenenrechte")
numbered([
    "Der Auftragsverarbeiter unterstützt den Verantwortlichen mit geeigneten technischen und "
    "organisatorischen Maßnahmen dabei, Anträge betroffener Personen auf Auskunft, Berichtigung, "
    "Löschung, Einschränkung, Datenübertragbarkeit und Widerspruch zu erfüllen "
    "(Art. 12 bis 23 DSGVO).",
    "Wendet sich eine betroffene Person unmittelbar an den Auftragsverarbeiter, leitet dieser das "
    "Anliegen unverzüglich an den Verantwortlichen weiter und beantwortet es nicht selbst.",
    "Stellt der Verantwortliche die Daten nicht selbst über die Anwendung bereit/löschen kann, "
    "führt der Auftragsverarbeiter entsprechende Maßnahmen nach Weisung aus.",
])

h("§ 8 Meldung von Verletzungen des Schutzes personenbezogener Daten")
numbered([
    "Der Auftragsverarbeiter informiert den Verantwortlichen unverzüglich, spätestens innerhalb "
    "von 24 Stunden nach Bekanntwerden, über Verletzungen des Schutzes personenbezogener Daten in "
    "seinem Verantwortungsbereich (Art. 33 Abs. 2 DSGVO).",
    "Die Meldung enthält mindestens: Art der Verletzung, betroffene Datenkategorien und ungefähre "
    "Zahl der Betroffenen, wahrscheinliche Folgen sowie ergriffene oder vorgeschlagene "
    "Gegenmaßnahmen.",
    "Der Auftragsverarbeiter unterstützt den Verantwortlichen bei dessen Melde- und "
    "Benachrichtigungspflichten gegenüber Aufsichtsbehörde und betroffenen Personen "
    "(Art. 33, 34 DSGVO).",
])

h("§ 9 Kontrollrechte des Verantwortlichen")
numbered([
    "Der Verantwortliche hat das Recht, die Einhaltung der Vorschriften über den Datenschutz und "
    "der vertraglichen Vereinbarungen im erforderlichen Umfang zu kontrollieren (Art. 28 Abs. 3 "
    "lit. h DSGVO).",
    "Der Auftragsverarbeiter weist die Einhaltung vorrangig durch geeignete Nachweise nach "
    "(Selbstauskunft, aktuelle Bescheinigungen, Zertifikate oder Prüfberichte, z. B. zur "
    "Infrastruktur des Hosting-Dienstleisters).",
    "Vor-Ort-Kontrollen werden mit angemessener Vorlaufzeit (i. d. R. 14 Tage) zu üblichen "
    "Geschäftszeiten und ohne Störung des Betriebsablaufs angekündigt und durchgeführt. "
    "Der Verantwortliche trägt die ihm hierdurch entstehenden Kosten.",
])

h("§ 10 Löschung und Rückgabe nach Beendigung")
numbered([
    "Nach Abschluss der Verarbeitungstätigkeiten löscht der Auftragsverarbeiter nach Wahl des "
    "Verantwortlichen alle personenbezogenen Daten oder gibt sie zurück und löscht vorhandene "
    "Kopien, sofern keine gesetzliche Aufbewahrungspflicht entgegensteht (Art. 28 Abs. 3 lit. g "
    "DSGVO).",
    "Auf Wunsch stellt der Auftragsverarbeiter die Daten in einem gängigen, maschinenlesbaren "
    "Format (Export) bereit.",
    "Die Löschung schließt Backups ein; diese werden im Rahmen des regulären Backup-Zyklus "
    "überschrieben. Der Auftragsverarbeiter bestätigt die Löschung auf Verlangen schriftlich.",
])

h("§ 11 Haftung")
body("Die Haftung der Parteien richtet sich nach Art. 82 DSGVO und den Regelungen des Hauptvertrags. "
     "Im Außenverhältnis gegenüber betroffenen Personen haften die Parteien nach den "
     "gesetzlichen Bestimmungen; im Innenverhältnis trägt jede Partei die Verantwortung für "
     "den von ihr zu vertretenden Anteil.")

h("§ 12 Schlussbestimmungen")
numbered([
    "Änderungen und Ergänzungen dieses Vertrags bedürfen der Textform. Dies gilt auch für "
    "die Aufhebung dieses Formerfordernisses.",
    "Bei Widersprüchen zwischen diesem Vertrag und dem Hauptvertrag gehen die Regelungen dieses "
    "Vertrags zum Datenschutz vor.",
    "Sollte eine Bestimmung unwirksam sein, bleibt die Wirksamkeit des übrigen Vertrags unberührt. "
    "An die Stelle der unwirksamen Bestimmung tritt eine Regelung, die dem wirtschaftlichen Zweck am "
    "nächsten kommt.",
    "Es gilt das Recht der Bundesrepublik Deutschland. Gerichtsstand ist, soweit zulässig, der "
    "Sitz des Auftragsverarbeiters.",
])

# ---- Unterschriften
h("Unterschriften")
body("Dieser Vertrag tritt mit Unterzeichnung beider Parteien in Kraft.", space_after=12)

sig = doc.add_table(rows=2, cols=2)
sig.alignment = WD_TABLE_ALIGNMENT.CENTER
c = sig.cell(0, 0).paragraphs[0]
c.add_run("Verantwortlicher (Kunde)").bold = True
c2 = sig.cell(0, 1).paragraphs[0]
c2.add_run("Auftragsverarbeiter (Anbieter)").bold = True
left = sig.cell(1, 0)
right = sig.cell(1, 1)
for cell in (left, right):
    cell.add_paragraph("\n\n____________________________________")
    cell.add_paragraph("Ort, Datum")
    cell.add_paragraph("\n\n____________________________________")
    cell.add_paragraph("Unterschrift, Name in Druckbuchstaben")
    cell.add_paragraph("Funktion: ___________________________")
for i in (0, 1):
    for row in sig.rows:
        row.cells[i].width = Mm(89)

# ============================================================ ANLAGE 1: TOM
doc.add_page_break()
title("Anlage 1: Technische und organisatorische Maßnahmen")
subtitle("gemäß Art. 32 DSGVO – Stand " + date.today().strftime("%d.%m.%Y"))

body("Die nachfolgenden Maßnahmen beschreiben das Schutzniveau des Auftragsverarbeiters für die "
     "im Auftrag verarbeiteten personenbezogenen Daten. Das Hosting erfolgt in Rechenzentren innerhalb "
     "Deutschlands/der EU.", space_after=6)

tom = [
    ("Vertraulichkeit – Zutrittskontrolle",
     "Hosting in zertifizierten Rechenzentren des Infrastruktur-Dienstleisters (Deutschland/EU) mit "
     "Zutrittssicherung (Zaun, Vereinzelung, Videoüberwachung, Zutrittsprotokollierung). Kein physischer "
     "Serverzugang durch den Auftragsverarbeiter selbst."),
    ("Vertraulichkeit – Zugangskontrolle",
     "Persönliche Benutzerkonten mit starken Passwörtern; Administrationszugänge nur über "
     "gesicherte Verbindungen (SSH-Key / 2-Faktor). Anwendungszugang über Lizenz-Gate und "
     "Benutzer-Login je Mandant; automatische Sperrung nach Inaktivität."),
    ("Vertraulichkeit – Zugriffskontrolle",
     "Rollen-/Rechtekonzept innerhalb der Anwendung (wer darf was). Strikte Mandantentrennung: jede "
     "Kundenfirma erhält eine eigene, getrennte Datenbank. Zugriff nur auf die jeweils eigenen Daten."),
    ("Vertraulichkeit – Trennungskontrolle",
     "Logische und physische Trennung der Mandantendaten (eine Datenbankdatei pro Firma). Trennung von "
     "Test-/Demo- und Produktivdaten."),
    ("Vertraulichkeit – Pseudonymisierung / Verschlüsselung",
     "Transportverschlüsselung aller Verbindungen per TLS/HTTPS (Zertifikate via Let's Encrypt). "
     "Verschlüsselung der Datenträger/Backups im Ruhezustand. Passwörter werden nur als "
     "Hash gespeichert."),
    ("Integrität – Weitergabekontrolle",
     "Keine Übermittlung an Dritte außerhalb der genehmigten Unterauftragsverarbeiter. "
     "Datentransfers ausschließlich verschlüsselt. Keine Datenübermittlung in Drittländer."),
    ("Integrität – Eingabekontrolle",
     "Protokollierung sicherheitsrelevanter Vorgänge (Anmeldungen, Änderungen, Exporte) mit "
     "Benutzer- und Zeitstempel, soweit fachlich vorgesehen (Audit-Log)."),
    ("Verfügbarkeit / Belastbarkeit",
     "Regelmäßige, automatisierte Backups; getestete Wiederherstellung. Überwachung der "
     "Erreichbarkeit (Uptime-Monitoring). Aktuelle Sicherheitsupdates des Betriebssystems und der "
     "Anwendung."),
    ("Wiederherstellbarkeit",
     "Definierter Wiederherstellungsprozess aus Backups; Backup-Aufbewahrung mehrerer Generationen. "
     "Ziel-Wiederherstellzeit und -punkt nach Stand der Technik für ein Produkt dieser Größe."),
    ("Verfahren zur regelmäßigen Überprüfung (Art. 32 Abs. 1 lit. d)",
     "Regelmäßige Überprüfung und Aktualisierung der Maßnahmen. Auftrags-/"
     "Weisungskontrolle durch dokumentierte Weisungen. Verpflichtung der Beschäftigten auf die "
     "Vertraulichkeit."),
]
t = doc.add_table(rows=0, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hr = t.add_row().cells
for i, txt in enumerate(("Schutzziel / Maßnahme", "Umsetzung")):
    set_cell_bg(hr[i], "0B6E6E")
    rr = hr[i].paragraphs[0].add_run(txt)
    rr.bold = True
    rr.font.size = Pt(9)
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for k, v in tom:
    cells = t.add_row().cells
    rk = cells[0].paragraphs[0].add_run(k)
    rk.bold = True
    rk.font.size = Pt(9)
    rv = cells[1].paragraphs[0].add_run(v)
    rv.font.size = Pt(9)
for i, w in enumerate((52, 126)):
    for row in t.rows:
        row.cells[i].width = Mm(w)
_set_table_borders(t)
doc.add_paragraph()
body("Hinweis: Die konkrete Ausprägung einzelner Maßnahmen (z. B. Backup-Intervalle, "
     "Aufbewahrungsdauer, 2-Faktor) ist vor Vertragsschluss an den tatsächlichen Betriebsstand "
     "anzupassen.", size=9, italic=True, color=GRAU)

# ============================================================ ANLAGE 2: SUB
doc.add_page_break()
title("Anlage 2: Genehmigte Unterauftragsverarbeiter")
subtitle("Stand " + date.today().strftime("%d.%m.%Y"))

body("Der Verantwortliche genehmigt den Einsatz der folgenden Unterauftragsverarbeiter "
     "(Art. 28 Abs. 2 DSGVO). Alle Anbieter verarbeiten innerhalb der EU/des EWR.", space_after=6)

subs = [
    ("Hetzner Online GmbH", "Industriestr. 25, 91710 Gunzenhausen, DE",
     "Hosting / Rechenzentrum (Server, Speicher, Backups)", "Deutschland"),
    ("Brevo (Sendinblue GmbH)", "Köpenicker Str. 126, 10179 Berlin, DE / FR",
     "Versand von Transaktions-E-Mails (Einladung, Passwort-Reset)", "EU"),
]
t = doc.add_table(rows=0, cols=4)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hr = t.add_row().cells
for i, txt in enumerate(("Anbieter", "Sitz / Anschrift", "Leistung", "Ort der Verarbeitung")):
    set_cell_bg(hr[i], "0B4A86")
    rr = hr[i].paragraphs[0].add_run(txt)
    rr.bold = True
    rr.font.size = Pt(8.5)
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for a, b, c, d in subs:
    cells = t.add_row().cells
    for j, val in enumerate((a, b, c, d)):
        rr = cells[j].paragraphs[0].add_run(val)
        rr.font.size = Pt(8.5)
        if j == 0:
            rr.bold = True
for i, w in enumerate((38, 50, 60, 30)):
    for row in t.rows:
        row.cells[i].width = Mm(w)
_set_table_borders(t)
doc.add_paragraph()

note_box([
    {"t": "Zahlungsdienstleister (gesonderte Rolle)", "bold": True, "size": 9, "color": BLAU},
    {"t": "Wird für die Abrechnung ein Zahlungsdienstleister (z. B. Stripe) eingesetzt, verarbeitet "
          "dieser Zahlungsdaten regelmäßig als eigenständig Verantwortlicher, nicht als "
          "Unterauftragsverarbeiter. Personalstammdaten werden an den Zahlungsdienstleister nicht "
          "übermittelt. Dies ist im Hauptvertrag/der Datenschutzerklärung gesondert darzustellen.",
     "size": 9},
], fill="EAF2FB")

body("\n", space_after=2)
body("Vorbehalt: Mustervorlage ohne Gewähr. Vor produktivem Einsatz anwaltlich prüfen lassen "
     "und die Platzhalter ([Anbieter-Firma], [Inhaber] usw.) ausfüllen. Stand: " +
     date.today().strftime("%d.%m.%Y") + ".", size=8, italic=True, color=GRAU)

# --- Speichern ----------------------------------------------------------------
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("OK ->", OUT)
